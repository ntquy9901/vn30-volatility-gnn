"""
Export GNN+HAR architecture diagram + explanation to Word (.docx).

Steps:
  1. Write mermaid source to temp .mmd file
  2. Render PNG with mmdc (mermaid-cli)
  3. Build Word document with python-docx

Output: results/gnn_har_architecture.docx

Usage:
  python scripts/eda/export_architecture_doc.py
"""
import sys
import subprocess
from pathlib import Path

sys.stdout.reconfigure(line_buffering=True)

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_root      = Path(__file__).parent.parent.parent
RESULTS    = _root / "results"
RESULTS.mkdir(exist_ok=True)
OUT_DOCX    = RESULTS / "gnn_har_architecture.docx"
DIAGRAM_PNG = RESULTS / "gnn_har_architecture_diagram.png"

# ─────────────────────────── MERMAID SOURCE ───────────────────────────────────
# SISO architecture: 30 nodes (VN30 only, no VNINDEX), 4 separate models
MERMAID_SRC = """flowchart TD
    subgraph INPUT["Du lieu dau vao"]
        A["close_prices\\n5579 ngay x 30 tickers (VN30 only)"]
        A --> B["log_returns\\n5579 x 30"]
    end

    subgraph GRAPH["Xay dung Graph (1 lan, tu train)"]
        B --> G1["Pearson correlation\\ntren toan bo train period"]
        G1 --> G2["threshold >= 0.4 OR cung nganh\\n-> edge ton tai (2 chieu)"]
        G2 --> G3["DGLGraph\\n30 nodes, undirected\\nNo VNINDEX hub"]
    end

    subgraph SNAP["Build Snapshots (SISO, per horizon h)"]
        B --> F1["rv_d = abs(log_return[t])"]
        B --> F2["rv_w = std(log_return[t-4..t])"]
        B --> F3["rv_m = std(log_return[t-19..t])"]
        F1 & F2 & F3 --> X["X: n_snaps x 30 x 3\\n[rv_d, rv_w, rv_m]"]
        A --> T1["compute_rv(close, h=h)\\n-> RV target"]
        T1 --> Y["y: n_snaps x 30\\nRV thuc te cho horizon h"]
    end

    subgraph HAR["HAR Baseline (OLS, fit tren train)"]
        Y --> R1["y_residual = y - HAR_pred\\nn_snaps x 30"]
        A --> HAR1["fit_har() -> coefficients\\npredict_har() -> HAR_pred"]
        HAR1 --> R1
    end

    subgraph NORM["Z-score Normalization (tu train)"]
        X --> XN["X_norm = (X - feat_mu) / feat_sig\\nshape: (30, 3)"]
        R1 --> YN["y_norm = (residual - rv_mu) / rv_sig\\nshape: (30,)"]
    end

    subgraph MODEL["GNNHARModel (SISO) - forward()"]
        XN --> C1["SAGEConv(3 -> 16, agg=mean)"]
        G3 -.->|graph structure| C1
        C1 --> E1["ELU"]
        E1 --> D1["Dropout(0.1)"]
        D1 --> C2["SAGEConv(16 -> 16, agg=mean)"]
        G3 -.->|graph structure| C2
        C2 --> E2["ELU"]
        E2 --> D2["Dropout(0.1)"]
        D2 --> H["hidden: 30 x 16"]
        H --> HEAD["Linear(16 -> 1)\\n.squeeze(-1)"]
        HEAD --> PRED["pred_norm: (30,)"]
    end

    subgraph LOSS["Training Loss (per model, per horizon)"]
        PRED --> MSE["MSE(pred_norm, y_norm)\\nAdamW + clip_grad_norm(1.0)"]
        YN --> MSE
    end

    subgraph DENORM["Denormalization + Reassembly"]
        PRED --> DN["residual = pred_norm * rv_sig + rv_mu"]
        HAR1 --> FP["final_pred = clip(HAR_pred + residual, min=0)"]
        DN --> FP
    end

    subgraph EVAL["Evaluation (per stock, per horizon)"]
        FP --> M1["R2"]
        FP --> M2["MAE"]
        FP --> M3["RMSE"]
        FP --> M4["QLIKE"]
    end

    style MODEL fill:#e8f4f8,stroke:#2196F3
    style HAR fill:#fff3e0,stroke:#FF9800
    style LOSS fill:#fce4ec,stroke:#E91E63
    style DENORM fill:#e8f5e9,stroke:#4CAF50
"""


# ─────────────────────────── RENDER MERMAID ───────────────────────────────────
def render_mermaid(src: str, out_png: Path) -> bool:
    import os
    npm_bin  = Path(os.environ.get("APPDATA", "")) / "npm"
    mmdc_cmd = str(npm_bin / "mmdc.cmd") if (npm_bin / "mmdc.cmd").exists() else "mmdc"

    mmd_file = RESULTS / "_temp_arch.mmd"
    mmd_file.write_text(src, encoding="utf-8")
    try:
        result = subprocess.run(
            [mmdc_cmd, "-i", str(mmd_file), "-o", str(out_png),
             "-w", "1400", "-H", "900", "-b", "white"],
            capture_output=True, text=True, timeout=60,
        )
        mmd_file.unlink(missing_ok=True)
        if result.returncode == 0 and out_png.exists():
            print(f"  [OK] Diagram rendered: {out_png.name}  ({out_png.stat().st_size//1024} KB)")
            return True
        print(f"  [WARN] mmdc failed: {result.stderr[:200]}")
        return False
    except Exception as e:
        print(f"  [WARN] mmdc error: {e}")
        mmd_file.unlink(missing_ok=True)
        return False


# ─────────────────────────── WORD HELPERS ─────────────────────────────────────
def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.size = Pt(11)


def add_code(doc: Document, text: str):
    p = doc.add_paragraph(text)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    if p.runs:
        p.runs[0].font.size = Pt(11)


def add_table_row(table, cells: list, bold_first=False):
    row = table.add_row()
    for i, val in enumerate(cells):
        cell = row.cells[i]
        cell.text = val
        if bold_first and i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True


# ─────────────────────────── BUILD DOCUMENT ───────────────────────────────────
def build_doc(diagram_ok: bool):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title ──────────────────────────────────────────────────────────────────
    title = doc.add_heading(
        "Kien truc GNN+HAR SISO — VN30 Realized Volatility Forecasting", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Mo hinh: GNNHARModel (GraphSAGE + SISO) voi HAR residual training | "
        "4 model doc lap, moi model cho 1 horizon h in [1, 5, 10, 20]"
    )
    doc.add_paragraph(
        "Tap du lieu: VN30 (30 co phieu), daily OHLCV, 2006-2026 | "
        "Test: 2026-01-01 -> 2026-04-14")
    doc.add_paragraph()

    # ── Diagram ────────────────────────────────────────────────────────────────
    add_heading(doc, "1. So do kien truc tong the", 1)
    if diagram_ok and DIAGRAM_PNG.exists():
        doc.add_picture(str(DIAGRAM_PNG), width=Inches(6.0))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(
            "Hinh 1: Kien truc GNN+HAR SISO — 30 nodes (VN30 only), "
            "4 model rieng biet moi horizon")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
    else:
        add_body(doc, "[Diagram PNG khong co san — chay mmdc de render]")
        add_code(doc, MERMAID_SRC[:400] + "\n... (xem file .mmd de biet them)")
    doc.add_paragraph()

    # ── Section 2: Tong quan ──────────────────────────────────────────────────
    add_heading(doc, "2. Tong quan kien truc (SISO)", 1)
    add_body(doc,
        "GNN+HAR ket hop mo hinh tuyen tinh HAR-RV voi mang Graph Neural Network "
        "theo cau truc skip connection. GNN khong hoc du bao RV tu dau ma chi hoc "
        "phan sai so cua HAR (HAR residuals). Ket qua cuoi cung la:"
    )
    add_code(doc, "final_pred = HAR_pred + GNN_correction")
    add_body(doc,
        "Neu GNN du bao correction = 0, thi final_pred = HAR_pred. "
        "Mo hinh co san la khong the kem hon HAR-RV (guaranteed performance floor).\n\n"
        "SISO (Single-Input Single-Output): 4 model doc lap, moi model du bao 1 horizon. "
        "Ly do tach SISO thay vi MIMO: R2 cua h=1 va h=20 khac xa nhau (-0.07 vs +0.90), "
        "trung binh loss nhieu horizon lam lech gradient cua nhau (gradient interference)."
    )
    doc.add_paragraph()

    # ── Section 3: Cac khoi ───────────────────────────────────────────────────
    add_heading(doc, "3. Giai thich tung khoi", 1)

    # --- 3.1
    add_heading(doc, "3.1 Du lieu dau vao (INPUT)", 2)
    add_body(doc,
        "Du lieu la gia dong hang ngay cua 30 co phieu VN30 (khong co VNINDEX), "
        "khoang 5579 ngay (2006-2026). Log-return duoc tinh bang:"
    )
    add_code(doc, "log_return[t] = ln(price[t] / price[t-1])")
    add_body(doc,
        "Log-return duoc su dung thay cho price return vi phan phoi xap xi normal hon, "
        "thich hop hon cho cac mo hinh thong ke va neural."
    )

    # --- 3.2
    add_heading(doc, "3.2 Xay dung Graph (30 nodes, khong co VNINDEX)", 2)
    add_body(doc,
        "1 graph tinh duoc xay dung 1 lan tu toan bo period train. "
        "Hai co phieu duoc noi canh neu tuong quan Pearson >= 0.4 HOAC cung nganh. "
        "Canh la vo huong (2 chieu). Node bi co lap duoc them self-loop."
    )
    add_bullet(doc, "30 nodes: VN30_TICKERS[0..29] (khong co VNINDEX)")
    add_bullet(doc, "Edges: undirected, Pearson(train) >= 0.4 OR same_sector")
    add_bullet(doc, "Graph tinh: khong thay doi trong suot qua trinh training va test")
    add_bullet(doc, "Tai sao bo VNINDEX: SAGEConv aggregate tu in-neighbors; "
               "canh VNINDEX->stock mot chieu nen VNINDEX khong nhan thong tin tu cac stock (loi thiet ke)")

    # --- 3.3
    add_heading(doc, "3.3 Build Snapshots (SISO, cho tung horizon h)", 2)
    add_body(doc, "Moi snapshot tai ngay t bao gom 3 features (cung loai do luong voi target):")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Feature"; hdr[1].text = "Cong thuc"; hdr[2].text = "Y nghia"
    for h in hdr:
        for run in h.paragraphs[0].runs:
            run.bold = True
    add_table_row(tbl, ["rv_d[t]", "abs(log_return[t])", "Bien dong 1 ngay (h=1 proxy)"])
    add_table_row(tbl, ["rv_w[t]", "std(log_return[t-4..t])", "Bien dong 5 ngay"])
    add_table_row(tbl, ["rv_m[t]", "std(log_return[t-19..t])", "Bien dong 20 ngay"])
    doc.add_paragraph()
    add_body(doc,
        "Target y[t] = std(log_return[t+1..t+h]) cho horizon h. "
        "Khong co leakage: feature dung [t-k..t], target dung [t+1..t+h]. "
        "Stride = MAX_H = 20 cho train/val (R4 compliance), stride = 1 cho test.\n"
        "X shape: (n_snaps, 30, 3) | y shape: (n_snaps, 30)"
    )

    # --- 3.4
    add_heading(doc, "3.4 HAR Baseline — OLS Residual Training (Core Innovation)", 2)
    add_body(doc,
        "Day la diem cot loi phan biet voi paper GNNHAR1L. "
        "HAR (Heterogeneous AutoRegressive) la mo hinh tuyen tinh OLS kinh dien:"
    )
    add_code(doc, "RV[t+h] = b0 + b1*RV[t] + b2*RV_w[t] + b3*RV_m[t] + epsilon")
    add_body(doc,
        "Coefficients b0..b3 duoc fit bang OLS tren tap train (khong refit tren val/test). "
        "GNN duoc train de du bao phan du epsilon (y_residual = y - HAR_pred), "
        "khong phai RV truc tiep. "
        "Neu GNN_correction = 0, final_pred = HAR_pred — mo hinh khong the kem hon HAR."
    )

    # --- 3.5
    add_heading(doc, "3.5 Z-score Normalization (SISO — per stock)", 2)
    add_code(doc,
        "X_norm = (X - feat_mu) / feat_sig    # shape: (30, 3) -- per node x per feature\n"
        "y_norm = (residual - rv_mu) / rv_sig  # shape: (30,)  -- per stock\n"
        "# rv_mu xap xi 0 (HAR la unbiased estimator)"
    )
    add_body(doc,
        "Normalize per-stock vi moi co phieu co scale bien dong khac nhau. "
        "Tat ca thong ke normalize duoc tinh tu TAP TRAIN ONLY va ap dung cho val/test."
    )

    # --- 3.6
    add_heading(doc, "3.6 GNNHARModel (SISO) — Forward Pass", 2)
    add_body(doc,
        "Kien truc 2-layer GraphSAGE voi 1 output head duy nhat (SISO). "
        "657 tham so: SAGEConv(3->16)=112, SAGEConv(16->16)=528, Linear(16->1)+bias=17."
    )

    steps = [
        ("SAGEConv(3 -> 16, agg=mean)",
         "Moi node tong hop thong tin hang xom bang trung binh, concat voi feature chinh no, "
         "nhan ma tran trong so W. Vi du: VCB sau buoc nay biet them ve ACB, TCB, MBB."),
        ("ELU activation",
         "ELU(x) = x neu x >= 0, else alpha*(e^x - 1). "
         "Khac ReLU: ELU giu gradient vung am, tranh dead neuron khi features z-scored co gia tri am."),
        ("Dropout(0.1)",
         "Tat ngau nhien 10% neurons khi training. hidden=16 nen chi dung 0.1 "
         "(dropout=0.3 se giet 5/16 neurons theo C5 — qua nhieu cho mo hinh nho)."),
        ("SAGEConv(16 -> 16, agg=mean)",
         "Lop thu 2 cho phep thong tin lan truyen qua 2 hops. "
         "Sau 2 lop, moi node 'nhin thay' vung lan can ban kinh 2."),
        ("Linear(16 -> 1) -> squeeze(-1) -> (30,)",
         "1 output head duy nhat. Khong co loss_mask, khong co MIMO cat. "
         "forward() tra ve (30,) truc tiep — 30 stocks, 1 horizon. "
         "Moi horizon chay mot model rieng biet."),
    ]
    for step_name, desc in steps:
        p = doc.add_paragraph()
        run_bold = p.add_run(f"  {step_name}: ")
        run_bold.bold = True
        run_bold.font.size = Pt(10)
        run_norm = p.add_run(desc)
        run_norm.font.size = Pt(10)

    # --- 3.7
    add_heading(doc, "3.7 Training Loss va Optimizer (per model, per horizon)", 2)
    add_code(doc,
        "# Model h=1 trains on:\n"
        "Loss = MSE(pred_norm_h1, y_norm_h1)    # shape: (30,)\n\n"
        "# Model h=5 trains on:\n"
        "Loss = MSE(pred_norm_h5, y_norm_h5)    # shape: (30,)\n\n"
        "Optimizer: AdamW(lr=1e-3, weight_decay=1e-4)\n"
        "Gradient clipping: clip_grad_norm(max_norm=1.0)\n"
        "Early stopping: patience=50 epochs | EPOCHS=500"
    )
    add_body(doc,
        "4 model doc lap, moi model optimize loss cua horizon rieng. "
        "clip_grad_norm(1.0) ngan exploding gradient khi co ngay COVID spike. "
        "SISO tranh gradient interference: h=1 va h=20 co noise level khac xa nhau."
    )

    # --- 3.8
    add_heading(doc, "3.8 Denormalization va Reassembly", 2)
    add_code(doc,
        "gnn_residual = pred_norm * rv_sig + rv_mu   # dua ve original residual scale\n"
        "final_pred   = HAR_pred + gnn_residual       # HAR + GNN correction\n"
        "final_pred   = clip(final_pred, min=0)       # RV >= 0 (std khong the am)"
    )
    add_body(doc,
        "GNN hoc correction trong z-score space; "
        "output cuoi can o don vi goc (RV = std of log_return). "
        "Clip(0) xu ly truong hop HAR_pred nho va GNN correction am qua."
    )

    # ── Section 4: So sanh voi paper ──────────────────────────────────────────
    add_heading(doc, "4. So sanh voi Paper GNNHAR1L (Zhang et al. IJF 2024)", 1)
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = "Table Grid"
    hdr2 = tbl2.rows[0].cells
    hdr2[0].text = ""; hdr2[1].text = "Paper GNNHAR1L"; hdr2[2].text = "Ours (GNN+HAR SISO)"
    for h in hdr2:
        for run in h.paragraphs[0].runs:
            run.bold = True
    rows2 = [
        ("H1 (skip)", "Trainable Linear(3,4) — hoc tu data", "Pre-computed OLS HAR"),
        ("H2 (GNN)", "GCN (GraphConv)", "SAGEConv 2-layer"),
        ("Graph", "N=? nodes", "30 nodes (VN30 only, no VNINDEX)"),
        ("Output", "relu(H1 + H2)", "clip(HAR + denorm(GNN), 0)"),
        ("Activation", "ReLU (gay mode collapse)", "ELU (xu ly am tot hon)"),
        ("Multi-horizon", "MIMO (1 model, 4 heads)", "SISO (4 models doc lap)"),
        ("Training", "Rolling retrain moi 22 ngay", "Static split (1 lan)"),
        ("Params", "Nhieu hon", "657 params"),
        ("R2 h=5 (test)", "-0.361 (Paper tren VN30)", "+0.676 (23/30 stocks > HAR)"),
        ("R2 h=10 (test)", "-0.719 (Paper tren VN30)", "+0.877 (28/30 stocks > HAR)"),
        ("R2 h=20 (test)", "-1.386 (Paper tren VN30)", "+0.933 (28/30 stocks > HAR)"),
        ("R2 h=1 (test)", "-0.125 (Paper tren VN30)", "-0.055 (HAR wins, noise floor)"),
    ]
    for r in rows2:
        add_table_row(tbl2, list(r), bold_first=True)
    doc.add_paragraph()

    # ── Section 5: Ket qua SISO (final) ─────────────────────────────────────
    add_heading(doc, "5. Ket qua thi nghiem (SISO clean, test 2026-01-01, 30 stocks)", 1)
    add_body(doc,
        "Ket qua duoi day tu kien truc SISO sach (30 nodes, khong co VNINDEX, 4 model doc lap). "
        "Test period: 2026-01-01 den 2026-04-14 | stride=1 (moi ngay giao dich)."
    )

    # Main results table
    tbl3 = doc.add_table(rows=1, cols=6)
    tbl3.style = "Table Grid"
    hdr3 = tbl3.rows[0].cells
    for i, h in enumerate(["Horizon", "n_test", "GNN R2", "HAR R2", "delta R2", "GNN > HAR"]):
        hdr3[i].text = h
        for run in hdr3[i].paragraphs[0].runs:
            run.bold = True
    results_data = [
        ("h=1",  "86",  "-0.055", "-0.031", "-0.023", "6/30"),
        ("h=5",  "82",  "+0.676", "+0.631", "+0.046", "23/30"),
        ("h=10", "77",  "+0.877", "+0.836", "+0.041", "28/30"),
        ("h=20", "67",  "+0.933", "+0.900", "+0.033", "28/30"),
    ]
    for r in results_data:
        add_table_row(tbl3, list(r))
    doc.add_paragraph()

    # QLIKE table
    add_body(doc, "QLIKE (Patton 2011) — chi tinh cho h=5,10,20 (h=1 bat thuong do RV gap 0):")
    tbl_q = doc.add_table(rows=1, cols=4)
    tbl_q.style = "Table Grid"
    hdr_q = tbl_q.rows[0].cells
    for i, h in enumerate(["Horizon", "GNN QLIKE", "HAR QLIKE", "delta"]):
        hdr_q[i].text = h
        for run in hdr_q[i].paragraphs[0].runs:
            run.bold = True
    qlike_data = [
        ("h=5",  "0.0528", "0.0620", "-0.0092 (GNN thap hon = tot hon)"),
        ("h=10", "0.0091", "0.0117", "-0.0026"),
        ("h=20", "0.0021", "0.0034", "-0.0013"),
    ]
    for r in qlike_data:
        add_table_row(tbl_q, list(r))
    doc.add_paragraph()

    add_body(doc,
        "Ket luan: OUTCOME #1 xac nhan — GNN+HAR SISO thang HAR-RV tai h=5 (23/30 co phieu), "
        "h=10 (28/30), h=20 (28/30). Cross-stock spillover qua GraphSAGE bo sung thong tin "
        "do luong duoc tai cac ky han dai, nhat quan voi Diebold & Yilmaz (2009)."
    )
    doc.add_paragraph()

    add_body(doc,
        "Nhan xet 1 — h=1 (GNN kem hon HAR, 6/30): "
        "Bien dong 1 ngay co muc nhieu rat cao. GNN aggregate thong tin tu hang xom "
        "nhung hang xom cung nhieu nhu nhau tai h=1. HAR marginally better voi per-stock linear. "
        "Day la ket qua binh thuong, xac nhan graph chỉ co gia tri tai ky han dai."
    )
    doc.add_paragraph()
    add_body(doc,
        "Nhan xet 2 — PLX (Petrolimex) thua ca 3 ky han dai h=5,10,20: "
        "PLX la co phieu nang luong/dau khi, bien dong theo gia dau quoc te — "
        "co cau truc khac biet so voi ngan hang va bat dong san chiem phan lon VN30. "
        "GraphSAGE truyen thong tin tu hang xom banking/realty den PLX gay nhieu hon giup ich."
    )
    doc.add_paragraph()
    add_body(doc,
        "Nhan xet 3 — Top GNN wins tai h=5 (delta R2 lon nhat): "
        "SHB (+0.144), STB (+0.140), SSB (+0.140), FPT (+0.121), PDR (+0.118). "
        "Ngan hang va cong nghe — nhung co phieu nay co tuong quan cheo nganh cao nhat "
        "va huong loi nhieu nhat tu co che truyen thong tin qua GraphSAGE."
    )

    doc.save(str(OUT_DOCX))
    print(f"  [OK] Word document saved: {OUT_DOCX}")


# ─────────────────────────── MAIN ─────────────────────────────────────────────
def main():
    print(f"\n{'='*60}")
    print(f"  Export GNN+HAR Architecture to Word (SISO, 30 nodes)")
    print(f"{'='*60}\n")

    print("[1] Rendering Mermaid diagram ...")
    diagram_ok = render_mermaid(MERMAID_SRC, DIAGRAM_PNG)

    print("[2] Building Word document ...")
    build_doc(diagram_ok)

    print(f"\n  Open: {OUT_DOCX}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()
